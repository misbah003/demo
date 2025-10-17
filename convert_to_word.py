"""
Convert Markdown documentation to Word document
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def add_heading(doc, text, level=1):
    """Add a heading with custom formatting"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with optional formatting"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def add_code_block(doc, code):
    """Add a code block with monospace font"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    return p

def add_table_from_markdown(doc, table_text):
    """Parse markdown table and add to document"""
    lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]
    if len(lines) < 2:
        return
    
    # Parse header
    headers = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
    
    # Parse rows (skip separator line)
    rows = []
    for line in lines[2:]:
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:
            rows.append(cells)
    
    if not rows:
        return
    
    # Create table
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Add rows
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < len(table.rows[i+1].cells):
                table.rows[i+1].cells[j].text = cell_text

def convert_markdown_to_word(md_file, docx_file):
    """Convert markdown file to Word document"""
    doc = Document()
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block = []
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                add_code_block(doc, '\n'.join(code_block))
                code_block = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block.append(line)
            i += 1
            continue
        
        # Handle tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            add_table_from_markdown(doc, '\n'.join(table_lines))
            table_lines = []
            in_table = False
        
        # Handle headings
        if line.startswith('# '):
            add_heading(doc, line[2:], level=1)
        elif line.startswith('## '):
            add_heading(doc, line[3:], level=2)
        elif line.startswith('### '):
            add_heading(doc, line[4:], level=3)
        elif line.startswith('#### '):
            add_heading(doc, line[5:], level=4)
        
        # Handle horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph('_' * 50)
        
        # Handle lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Code
            doc.add_paragraph(text, style='List Bullet')
        
        elif line.strip().startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            text = re.sub(r'^\d+\.\s+', '', line.strip())
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            doc.add_paragraph(text, style='List Number')
        
        # Handle regular paragraphs
        elif line.strip():
            # Remove markdown formatting
            text = line.strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Code
            
            # Skip emoji-only lines or decorative lines
            if not re.match(r'^[─┌┐└┘│├┤┬┴┼]+$', text):
                doc.add_paragraph(text)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✅ Word document created: {docx_file}")

if __name__ == "__main__":
    md_file = r"c:\Users\HomeLaptop\Downloads\navi-tax-35-main\ML_Tax_System_Documentation.md"
    docx_file = r"c:\Users\HomeLaptop\Downloads\navi-tax-35-main\ML_Tax_System_Documentation.docx"
    
    convert_markdown_to_word(md_file, docx_file)